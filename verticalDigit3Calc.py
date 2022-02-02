import os, sys
sys.path.append(os.path.dirname(os.path.abspath(os.path.dirname(__file__))))
from verticalDigitNCalc import CalcPractice
from digit3Subtract import digit3Subtract
from digit3Add import digit3Add
import docx
import random
import verticalAnswerForDigit3N2Calc

class verticalDigit3Calc(CalcPractice):
    fileName = "digit3Calc"
    
    def __init__(self, baseForPracticeDocxName, baseForAnswerDocxName):
        CalcPractice.__init__(self, baseForPracticeDocxName, baseForAnswerDocxName)

        
    def getProblem(self):
        flag = random.randrange(0,2) #0 -> +, 1 -> -
        if flag == 0:
            c = digit3Add()

        else:
            c = digit3Subtract()

        return c

    def fileWrite(self, file, c):
        file.write('\n' + c[0] + '\n' + c[1] + '\n')
        
    
    def docxWrite(self, wFile, c):
        wFile.add_paragraph(c[0])
        wFile.add_paragraph(c[1],style='표준밑줄')
        wFile.add_paragraph('')

    def answering(self): #txt로 저장 포함
        verticalAnswerForDigit3N2Calc.answering(self.fileName)
    
    def txt2Docx(self, wName):
        verticalAnswerForDigit3N2Calc.answerText2Docx(self.fileName, wName)


if __name__ == '__main__':
    asdf = verticalDigit3Calc('baseForAddNSubtract.docx', 'baseForVerticalAnswer.docx')
    print(asdf.fileName + " generator")
    i = int(input('생성 할 문제 수(1page = 30문제): '))
    asdf.digitNCalc(i)
